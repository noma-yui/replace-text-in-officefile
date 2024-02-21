import argparse
from docx import Document
import yaml


# At python-docx
# Document contains paragraphs, tables, sections
# Paragraph contains text
# table constains cell
# cell contains paragraphs, tables、text
# _Header contains paragraphs, tables
# _Footer contains paragraphs, tables
# This script does not process headers and footers.


def parse_var_yml(var_yml):
    # var_yml -> tuple
    with open(var_yml, "r", encoding='UTF-8') as stream:
        data = yaml.load(stream, Loader=yaml.FullLoader)

    vars_vals = tuple(("$"+ k + "$", v) for k, v in data.items() )
    return vars_vals


def replace_pragraph(paragraph, vars_vals):
    tmp_text = paragraph.text
    for var, val in vars_vals:
        tmp_text = tmp_text.replace(var, val)
    if tmp_text != paragraph.text:
        # get style from the first run in a paragraph
        tmp_run = paragraph.runs[0]
        tmp_font = tmp_run.font
        tmp_style = tmp_run.style

        # rewrite
        paragraph.text = tmp_text
        tmp2_font = paragraph.runs[0].font
        tmp2_font.name = tmp_font.name
        tmp2_font.size = tmp_font.size
        tmp2_font.bold = tmp_font.bold
        tmp2_font.italic = tmp_font.italic
        tmp2_font.underline = tmp_font.underline
        paragraph.runs[0].style = tmp_style
    return

def replace_table(table, vars_vals):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_pragraph(para, vars_vals)
            for tbl in cell.tables:
                replace_table(tbl, vars_vals)
    return


def subsituting(doc, vars_vals, out_file):
    # replace
    for para in doc.paragraphs:
        replace_pragraph(para, vars_vals= vars_vals)
    # replace
    for table in doc.tables:
        replace_table(table, vars_vals= vars_vals)
    # save
    doc.save(out_file)


def main(template_file, out_file, var_yml):
    # var_yml -> tuple
    vars_vals = parse_var_yml(var_yml)
    # docs...
    doc = Document(template_file)
    subsituting(doc, vars_vals, out_file)


if __name__ == "__main__":
    ########################
    # parse arguments
    parser = argparse.ArgumentParser(description='Docx replacer')

    # input data
    parser.add_argument('--input',
                        type = str,
                        required = True,
                        dest = "template_file",
                        default = "template_file",
                        help='template file to be replace.')
    parser.add_argument('--out',
                        type = str,
                        required = False,
                        dest = "out_file",
                        help='output file name.')
    parser.add_argument('--var',
                        type = str,
                        required = True,
                        dest = "var_yml",
                        help='YAML files for keys and values')

    # parse arguments
    args = parser.parse_args()

    ########################
    # parse arguments
    # copy
    template_file = args.template_file
    out_file = args.out_file
    var_yml = args.var_yml

    if out_file is None:
        out_file = template_file + ".sub.docx"


    main(template_file, out_file, var_yml)
